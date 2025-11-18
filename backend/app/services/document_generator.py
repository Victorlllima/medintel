"""
Medical document generation service (PDF/DOCX)
"""
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime, date
import os
import tempfile

from app.schemas.document import DocumentType, DocumentFormat


class DocumentGenerator:
    """Generate medical documents in PDF or DOCX format"""

    def generate(
        self,
        document_type: DocumentType,
        format: DocumentFormat,
        patient: dict,
        user: dict,
        consultation: dict,
        data: dict
    ) -> str:
        """
        Generate a medical document

        Args:
            document_type: Type of document to generate
            format: Output format (PDF or DOCX)
            patient: Patient data
            user: Doctor/user data
            consultation: Consultation data
            data: Document-specific data

        Returns:
            Path to generated file
        """
        if format == DocumentFormat.PDF:
            return self._generate_pdf(document_type, patient, user, consultation, data)
        else:
            return self._generate_docx(document_type, patient, user, consultation, data)

    def _generate_pdf(
        self,
        document_type: DocumentType,
        patient: dict,
        user: dict,
        consultation: dict,
        data: dict
    ) -> str:
        """Generate PDF document"""
        # Create temporary file
        fd, temp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        doc = SimpleDocTemplate(temp_path, pagesize=A4)
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#0066CC',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        story = []

        if document_type == DocumentType.MEDICAL_CERTIFICATE:
            story = self._build_medical_certificate_pdf(patient, user, data, styles, title_style)
        elif document_type == DocumentType.PRESCRIPTION:
            story = self._build_prescription_pdf(patient, user, data, styles, title_style)
        elif document_type == DocumentType.ATTENDANCE_DECLARATION:
            story = self._build_attendance_declaration_pdf(patient, user, data, styles, title_style)

        doc.build(story)

        return temp_path

    def _build_medical_certificate_pdf(self, patient, user, data, styles, title_style):
        """Build medical certificate content"""
        story = []

        # Title
        story.append(Paragraph("ATESTADO MÉDICO", title_style))
        story.append(Spacer(1, 1*cm))

        # Content
        content = f"""
        Atesto para os devidos fins que <b>{patient['name']}</b>,
        portador(a) do CPF <b>{patient['cpf']}</b>, nascido(a) em <b>{patient['date_of_birth']}</b>,
        esteve sob meus cuidados médicos e necessita de afastamento de suas atividades
        por um período de <b>{data['days_off']} dia(s)</b>, a partir de <b>{data['start_date']}</b>.
        """

        if data.get('icd_code'):
            content += f"\n\nCID-10: <b>{data['icd_code']}</b>"

        if data.get('observations'):
            content += f"\n\n{data['observations']}"

        story.append(Paragraph(content, styles['BodyText']))
        story.append(Spacer(1, 2*cm))

        # Signature
        today = datetime.now().strftime("%d/%m/%Y")
        signature = f"""
        <br/><br/>
        _____________________________________________<br/>
        <b>Dr(a). {user['full_name']}</b><br/>
        CRM: {user['crm']}<br/>
        {user.get('specialty', '')}<br/>
        {today}
        """

        story.append(Paragraph(signature, styles['BodyText']))

        return story

    def _build_prescription_pdf(self, patient, user, data, styles, title_style):
        """Build prescription content"""
        story = []

        story.append(Paragraph("RECEITA MÉDICA", title_style))
        story.append(Spacer(1, 0.5*cm))

        # Patient info
        patient_info = f"""
        <b>Paciente:</b> {patient['name']}<br/>
        <b>Data:</b> {datetime.now().strftime("%d/%m/%Y")}
        """
        story.append(Paragraph(patient_info, styles['BodyText']))
        story.append(Spacer(1, 1*cm))

        # Medications
        for med in data.get('medications', []):
            med_text = f"""
            <b>{med.get('name', 'Medicamento')}</b><br/>
            Dosagem: {med.get('dosage', '')}<br/>
            Frequência: {med.get('frequency', '')}<br/>
            Duração: {med.get('duration', '')}<br/>
            """
            story.append(Paragraph(med_text, styles['BodyText']))
            story.append(Spacer(1, 0.5*cm))

        if data.get('observations'):
            story.append(Spacer(1, 0.5*cm))
            story.append(Paragraph(f"<b>Observações:</b> {data['observations']}", styles['BodyText']))

        story.append(Spacer(1, 2*cm))

        # Signature
        signature = f"""
        _____________________________________________<br/>
        <b>Dr(a). {user['full_name']}</b><br/>
        CRM: {user['crm']}
        """
        story.append(Paragraph(signature, styles['BodyText']))

        return story

    def _build_attendance_declaration_pdf(self, patient, user, data, styles, title_style):
        """Build attendance declaration content"""
        story = []

        story.append(Paragraph("DECLARAÇÃO DE COMPARECIMENTO", title_style))
        story.append(Spacer(1, 1*cm))

        # Content
        content = f"""
        Declaro para os devidos fins que <b>{patient['name']}</b>,
        portador(a) do CPF <b>{patient['cpf']}</b>, compareceu à consulta médica
        no dia <b>{data['date']}</b>
        """

        if data.get('start_time') and data.get('end_time'):
            content += f" das <b>{data['start_time']}</b> às <b>{data['end_time']}</b>"

        content += "."

        if data.get('purpose'):
            content += f"\n\n<b>Motivo:</b> {data['purpose']}"

        story.append(Paragraph(content, styles['BodyText']))
        story.append(Spacer(1, 2*cm))

        # Signature
        today = datetime.now().strftime("%d/%m/%Y")
        signature = f"""
        _____________________________________________<br/>
        <b>Dr(a). {user['full_name']}</b><br/>
        CRM: {user['crm']}<br/>
        {today}
        """
        story.append(Paragraph(signature, styles['BodyText']))

        return story

    def _generate_docx(
        self,
        document_type: DocumentType,
        patient: dict,
        user: dict,
        consultation: dict,
        data: dict
    ) -> str:
        """Generate DOCX document"""
        # Create temporary file
        fd, temp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        if document_type == DocumentType.MEDICAL_CERTIFICATE:
            self._build_medical_certificate_docx(doc, patient, user, data)
        elif document_type == DocumentType.PRESCRIPTION:
            self._build_prescription_docx(doc, patient, user, data)
        elif document_type == DocumentType.ATTENDANCE_DECLARATION:
            self._build_attendance_declaration_docx(doc, patient, user, data)

        doc.save(temp_path)

        return temp_path

    def _build_medical_certificate_docx(self, doc, patient, user, data):
        """Build medical certificate DOCX"""
        # Title
        title = doc.add_heading('ATESTADO MÉDICO', level=1)
        title.alignment = 1  # Center

        doc.add_paragraph()

        # Content
        p = doc.add_paragraph()
        p.add_run(f"Atesto para os devidos fins que {patient['name']}, ")
        p.add_run(f"portador(a) do CPF {patient['cpf']}, nascido(a) em {patient['date_of_birth']}, ")
        p.add_run(f"esteve sob meus cuidados médicos e necessita de afastamento de suas atividades ")
        p.add_run(f"por um período de {data['days_off']} dia(s), a partir de {data['start_date']}.")

        if data.get('icd_code'):
            doc.add_paragraph(f"CID-10: {data['icd_code']}")

        if data.get('observations'):
            doc.add_paragraph(data['observations'])

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()

        sig = doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"Dr(a). {user['full_name']}")
        doc.add_paragraph(f"CRM: {user['crm']}")
        if user.get('specialty'):
            doc.add_paragraph(user['specialty'])
        doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))

    def _build_prescription_docx(self, doc, patient, user, data):
        """Build prescription DOCX"""
        title = doc.add_heading('RECEITA MÉDICA', level=1)
        title.alignment = 1

        doc.add_paragraph(f"Paciente: {patient['name']}")
        doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")

        doc.add_paragraph()

        for med in data.get('medications', []):
            doc.add_paragraph(med.get('name', 'Medicamento'), style='Heading3')
            doc.add_paragraph(f"Dosagem: {med.get('dosage', '')}")
            doc.add_paragraph(f"Frequência: {med.get('frequency', '')}")
            doc.add_paragraph(f"Duração: {med.get('duration', '')}")
            doc.add_paragraph()

        if data.get('observations'):
            doc.add_paragraph(f"Observações: {data['observations']}")

        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"Dr(a). {user['full_name']}")
        doc.add_paragraph(f"CRM: {user['crm']}")

    def _build_attendance_declaration_docx(self, doc, patient, user, data):
        """Build attendance declaration DOCX"""
        title = doc.add_heading('DECLARAÇÃO DE COMPARECIMENTO', level=1)
        title.alignment = 1

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"Declaro para os devidos fins que {patient['name']}, ")
        p.add_run(f"portador(a) do CPF {patient['cpf']}, compareceu à consulta médica ")
        p.add_run(f"no dia {data['date']}")

        if data.get('start_time') and data.get('end_time'):
            p.add_run(f" das {data['start_time']} às {data['end_time']}")

        p.add_run(".")

        if data.get('purpose'):
            doc.add_paragraph(f"Motivo: {data['purpose']}")

        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"Dr(a). {user['full_name']}")
        doc.add_paragraph(f"CRM: {user['crm']}")
        doc.add_paragraph(datetime.now().strftime("%d/%m/%Y"))
